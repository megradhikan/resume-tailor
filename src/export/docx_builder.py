"""
DOCX export — builds a Word document from a resume with accepted
rewrite suggestions applied, or from a cover letter draft.
"""
from __future__ import annotations

import io
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FUZZY_THRESHOLD = 0.75  # similarity ratio to count as a match


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _find_replacement(line: str, replacements: dict[str, str]) -> str | None:
    """
    Look up a replacement for `line`.
    1. Exact match first (fast path).
    2. Fuzzy match: accept the best candidate above _FUZZY_THRESHOLD.
       Needed because the LLM often normalises whitespace or punctuation
       when quoting the original line back.
    """
    # Fast path — exact
    if line in replacements:
        return replacements[line]

    # Fuzzy path
    best_ratio = 0.0
    best_replacement: str | None = None
    for original, suggested in replacements.items():
        ratio = SequenceMatcher(None, line.lower(), original.lower()).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_replacement = suggested

    if best_ratio >= _FUZZY_THRESHOLD:
        return best_replacement
    return None


def build_resume_docx(
    resume_text: str,
    accepted_suggestions: list[dict],
) -> bytes:
    """
    Apply accepted rewrite suggestions to the resume text and export as DOCX.
    Changed lines are rendered in green so the user can review them.
    """
    replacements: dict[str, str] = {
        s["original_line"].strip(): s["suggested_line"].strip()
        for s in accepted_suggestions
    }

    doc = Document()
    doc.core_properties.title = "Tailored Resume"

    section_keywords = {
        "experience", "education", "skills", "projects",
        "summary", "certifications", "other",
    }

    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        replacement = _find_replacement(line, replacements)

        if replacement:
            p = doc.add_paragraph()
            run = p.add_run(replacement)
            run.font.color.rgb = RGBColor(0x16, 0x7A, 0x3E)
            run.bold = True
        elif len(line) < 60 and any(kw in line.lower() for kw in section_keywords):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    # Legend
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: green lines are accepted rewrite suggestions.")
    run.font.color.rgb = RGBColor(0x16, 0x7A, 0x3E)
    run.italic = True
    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cover_letter_docx(
    cover_letter_draft: str,
    company_name: str = "",
    role_title: str = "",
) -> bytes:
    """Export a cover letter draft as a clean DOCX."""
    doc = Document()
    doc.core_properties.title = f"Cover Letter — {role_title} at {company_name}".strip(" —")

    if company_name or role_title:
        title = f"{role_title} at {company_name}".strip(" at")
        doc.add_heading(title, level=1)
        doc.add_paragraph()

    for para in cover_letter_draft.split("\n\n"):
        text = para.strip()
        if text:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
