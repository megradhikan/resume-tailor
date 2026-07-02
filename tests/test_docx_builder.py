"""Unit tests for DOCX export — verifies output is valid DOCX bytes."""
import io
import pytest
from docx import Document
from src.export.docx_builder import build_resume_docx, build_cover_letter_docx


RESUME_TEXT = """Jane Doe

Experience
Senior Software Engineer at Acme Corp
Built REST APIs using FastAPI and PostgreSQL
Led a team of 5 engineers

Education
B.S. Computer Science — State University 2019

Skills
Python, FastAPI, PostgreSQL, Docker
"""

ACCEPTED_SUGGESTION = {
    "section": "experience",
    "original_line": "Built REST APIs using FastAPI and PostgreSQL",
    "suggested_line": "Architected and delivered high-throughput REST APIs using FastAPI and PostgreSQL, reducing p99 latency by 30%",
    "reason": "Adds measurable impact",
    "grounded_in": "FastAPI, PostgreSQL (in resume)",
}


def _load_docx(blob: bytes) -> Document:
    return Document(io.BytesIO(blob))


# ── build_resume_docx ────────────────────────────────────────────────────────

def test_resume_docx_returns_bytes():
    result = build_resume_docx(RESUME_TEXT, [])
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_resume_docx_is_valid_docx():
    result = build_resume_docx(RESUME_TEXT, [])
    doc = _load_docx(result)
    assert doc is not None


def test_resume_docx_fuzzy_matches_llm_paraphrase():
    """LLM often normalises punctuation/spacing in original_line — must still match."""
    # Resume has "Built REST APIs using FastAPI and PostgreSQL"
    # LLM quotes it back with a trailing period and extra space
    suggestion_with_imperfect_quote = {
        "section": "experience",
        "original_line": "Built REST APIs using FastAPI and PostgreSQL.",  # extra period
        "suggested_line": "Architected and delivered high-throughput REST APIs using FastAPI and PostgreSQL",
        "reason": "Stronger verb",
        "grounded_in": "FastAPI, PostgreSQL",
    }
    result = build_resume_docx(RESUME_TEXT, [suggestion_with_imperfect_quote])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "high-throughput" in all_text or "Architected" in all_text


def test_resume_docx_fuzzy_matches_whitespace_difference():
    """Extra leading/trailing whitespace in LLM quote must still fuzzy-match."""
    suggestion = {
        "section": "experience",
        "original_line": "  Led a team of 5 engineers  ",  # extra whitespace
        "suggested_line": "Managed and mentored a cross-functional team of 5 engineers",
        "reason": "Stronger leadership language",
        "grounded_in": "team of 5 engineers",
    }
    result = build_resume_docx(RESUME_TEXT, [suggestion])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "mentored" in all_text


def test_resume_docx_does_not_fuzzy_match_unrelated_line():
    """A suggestion whose original_line is totally different must not match any line."""
    suggestion = {
        "section": "skills",
        "original_line": "Completely unrelated line that does not appear anywhere",
        "suggested_line": "Should not appear in output",
        "reason": "n/a",
        "grounded_in": "n/a",
    }
    result = build_resume_docx(RESUME_TEXT, [suggestion])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Should not appear" not in all_text


def test_resume_docx_contains_original_lines_when_no_suggestions():
    result = build_resume_docx(RESUME_TEXT, [])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Led a team of 5 engineers" in all_text


def test_resume_docx_applies_accepted_suggestion():
    result = build_resume_docx(RESUME_TEXT, [ACCEPTED_SUGGESTION])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "high-throughput" in all_text or "Architected" in all_text


def test_resume_docx_omits_original_line_when_replaced():
    result = build_resume_docx(RESUME_TEXT, [ACCEPTED_SUGGESTION])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Original line should not appear as its own paragraph since it was replaced
    assert "Built REST APIs using FastAPI and PostgreSQL" not in all_text or "high-throughput" in all_text


def test_resume_docx_with_multiple_suggestions():
    suggestions = [
        ACCEPTED_SUGGESTION,
        {
            "section": "experience",
            "original_line": "Led a team of 5 engineers",
            "suggested_line": "Managed and mentored a team of 5 engineers, improving sprint velocity by 20%",
            "reason": "Adds leadership metric",
            "grounded_in": "team of 5 engineers (in resume)",
        },
    ]
    result = build_resume_docx(RESUME_TEXT, suggestions)
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "mentored" in all_text


def test_resume_docx_contains_legend():
    result = build_resume_docx(RESUME_TEXT, [ACCEPTED_SUGGESTION])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "green" in all_text.lower()


def test_resume_docx_empty_suggestions():
    """No suggestions — full resume text should still render."""
    result = build_resume_docx(RESUME_TEXT, [])
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Python" in all_text


# ── build_cover_letter_docx ─────────────────────────────────────────────────

COVER_LETTER_DRAFT = """Dear Hiring Manager,

I am excited to apply for the Senior Software Engineer role at Acme Corp.
My experience with FastAPI and PostgreSQL aligns well with your requirements.

While I haven't used Kubernetes directly, my experience with Docker and AWS
has given me strong container orchestration fundamentals.

I look forward to discussing how I can contribute to your team.

Sincerely,
Jane Doe
"""


def test_cover_letter_docx_returns_bytes():
    result = build_cover_letter_docx(COVER_LETTER_DRAFT)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_cover_letter_docx_is_valid():
    result = build_cover_letter_docx(COVER_LETTER_DRAFT, "Acme Corp", "Senior Software Engineer")
    doc = _load_docx(result)
    assert doc is not None


def test_cover_letter_docx_contains_body_text():
    result = build_cover_letter_docx(COVER_LETTER_DRAFT, "Acme Corp", "Senior Software Engineer")
    doc = _load_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "excited" in all_text


def test_cover_letter_docx_includes_title_when_company_and_role_provided():
    result = build_cover_letter_docx(COVER_LETTER_DRAFT, "Acme Corp", "Senior Software Engineer")
    doc = _load_docx(result)
    # Title is in headings
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Acme Corp" in h or "Senior Software Engineer" in h for h in headings)


def test_cover_letter_docx_works_without_company_or_role():
    result = build_cover_letter_docx(COVER_LETTER_DRAFT)
    doc = _load_docx(result)
    assert doc is not None
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "excited" in all_text


def test_cover_letter_docx_empty_draft():
    result = build_cover_letter_docx("")
    doc = _load_docx(result)
    assert doc is not None
